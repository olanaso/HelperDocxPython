# coding=utf-8
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
import os
import re

def buscar_add_parrafo(doc, textSearchParaffo,text_parrafo):
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Arial'
    paragraph=""
    for p in doc.paragraphs:
        if textSearchParaffo.decode('utf-8') in p.text:
            paragraph=p
            #p.alignment = 1
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text_parrafo:
        new_para.add_run(text_parrafo.decode('utf-8'))
    if style is not None:
        new_para.style = style
    return new_para

def buscar_add_parrafo3(doc,parrafo, text_parrafo):
    text_parrafo=text_parrafo.decode('utf-8')
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Arial'
    paragraph=parrafo
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text_parrafo:
        new_para.add_run(text_parrafo)
    if style is not None:
        new_para.style = style
    return new_para

def buscar_add_parrafo2(doc, textSearchParaffo,text_parrafo):
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Arial'
    paragraph=""
    for p in doc.paragraphs:
        if textSearchParaffo.decode('utf-8') in p.text:
            paragraph=p
    print(paragraph.text)
    paragraph.insert_paragraph_before(text_parrafo.decode('utf-8'))
    return doc

def replace_parrafo_parcial(doc, textSearchParaffo,textobuscar, text_remplazar):
    textobuscar=textobuscar.decode('utf-8')
    text_remplazar=text_remplazar.decode('utf-8')
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Arial'
    for p in doc.paragraphs:
        if textSearchParaffo.decode('utf-8') in p.text:
            inline = p.runs
            # print(inline)
            # Loop added to work with runs (strings with same style)
            '''for i in range(len(inline)):
                if 'old text' in inline[i].text:
                    text = inline[i].text.replace('old text', 'new text')
                    inline[i].text = text
            '''
            #p.encode('utf-8')
            p.text = p.text.replace(textobuscar, text_remplazar);
            p.style = doc.styles['Normal']
            #p.alignment = 1
    return doc

def replace_parrafo(doc, textSearchParaffo,text_parrafo):
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Arial'
    for p in doc.paragraphs:
        if textSearchParaffo.decode('utf-8') in p.text:
            inline = p.runs
            # print(inline)
            # Loop added to work with runs (strings with same style)
            '''for i in range(len(inline)):
                if 'old text' in inline[i].text:
                    text = inline[i].text.replace('old text', 'new text')
                    inline[i].text = text
            '''
            #p.encode('utf-8')
            p.text = text_parrafo.decode('utf-8');
            p.style = doc.styles['Normal']
            #p.alignment = 1
    return doc

def replace_parrafo_text(doc, textSearchParaffo,text_parrafo,textoantiguo,textonuevo):
    textoantiguo=textoantiguo.decode('utf-8')
    textonuevo=textonuevo.decode('utf-8')
    textSearchParaffo=textSearchParaffo.decode('utf-8')
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)
    font.name = 'Arial'
    for p in doc.paragraphs:
        if textSearchParaffo in p.text:
            inline = p.runs
            print(inline)
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if textoantiguo in inline[i].text:
                    text = inline[i].text.replace(textoantiguo, textonuevo)
                    inline[i].text = text

            #p.encode('utf-8')
            print(p.text)
            p.text = text_parrafo.decode('utf-8');
            p.style = doc.styles['Normal']
            #p.alignment = 1

    return doc

def buscar_eliminar(doc, textSearchParaffo):
    for p in doc.paragraphs:
        if textSearchParaffo.decode('utf-8') in p.text:
            p.clear()
            #p.alignment = 1
    return doc

def buscar_parrafo_parcial(doc, textSearchParaffo):
    parrafobuscado=""
    for p in doc.paragraphs:
        if textSearchParaffo.decode('utf-8') in p.text:
            parrafobuscado=p.text
    return parrafobuscado


rutaTdrs="D:\\GTS\TDRS\\abril\\test\\"
rutaTdrscambiados="D:\\GTS\TDRS\\abril\\test\\pro\\"

#obteniendo los archivos word de cada uno
files = [os.path.join(r,file) for r,d,f in os.walk(rutaTdrs) for file in f]
arr_txt = [x for x in files if x.endswith(".docx")]

for file in arr_txt:
    x = file.split("\\") #archivos separados
    print(file.decode('utf-8'))
    document = Document(file)

    # 1. Duracion del servicio
    replace_parrafo_parcial(document, 'El servicio se ejecutará en el plazo de hasta',
                            'contados a partir de la fecha de la suscripción del acta señalada en el numeral VI del presente documento',
                            "contados a partir del día de la notificación de la orden de servicio")

    # 2. Requisitos del proveedor
    # Segundo caso
    replace_parrafo_parcial(document, 'Contar con seguro complementario de trabajo y riesgo',
                            'Contar con seguro complementario de trabajo y riesgo (SCTR) en salud y pensión, o el que haga sus veces o supere, en cuánto a la cobertura de salud y pensión',
                            "Contar con seguro complementario de trabajo y riesgo (SCTR) de salud y pensión, por el tiempo de ejecución del servicio, el cual será entregada para la emisión de la orden de servicio")
    # teer caso
    buscar_add_parrafo2(document, 'Condiciones Específicas',
                        "Nota: La póliza de Seguro Complementario de Trabajo de Riesgo (SCTR) deberá ser acreditado una vez adjudicado el servicio, para la emisión de la respetiva orden de servicio, por el tiempo de ejecución del servicio.")


    # 3. Modificando el LUGAR DE SERVICIO
    searchbox=buscar_parrafo_parcial(document,"Contratar un (1) servicio ")
    pattern = "(?<= de C)(.*)(?=en el marco)|(?<=del )(.*)(?=en el marco)|(?<=de los )(.*)(?=en el marco)|(?<=correspondiente a la )(.*)(?=en el marco)|(?<=de las )(.*)(?=en el marco)|(?<=en las )(.*)(?=en el marco)"
    searchbox_result=""
    try:
        searchbox_result = re.search(pattern, searchbox).group()
    except AttributeError:
        searchbox_result = re.search(pattern, searchbox)

    searchbox_result=unicode(searchbox_result)
    #valorecontrado=str(searchbox_result).encode('utf-8')
    #valorecontrado=(''.join(searchbox_result)).encode('utf-8')
    #searchbox_result=str(searchbox_result)
    #valorecontrado= searchbox_result.encode('utf-8', 'ignore').decode('ascii')
    print (type (searchbox_result))

    buscar_eliminar(document, 'La prestación del servicio se ejecutará en el departamento');
    reemplazante = "La prestación del servicio se ejecutará en " + searchbox_result.encode('utf-8')

    print (type(reemplazante))

    print (reemplazante)
   # print(type(reemplazante))
    buscar_add_parrafo2(document, "FORMA DE PAGO",reemplazante)
    #m = re.search(r"(?<=departamento )(.*)(?=en el marco)",s)

    # 4. Penalidades
    buscar_eliminar(document,'Si el contratista incurre en retraso injustificado a la presentación')
    buscar_eliminar(document,'La acumulación del monto')
    parrafo_agregado=buscar_add_parrafo(document,"PENALIDADES","En caso de retraso en la ejecución de las prestaciones, se aplicará una penalidad al contratista por cada día de retraso hasta por el monto máximo del 10% del monto del contrato, según lo dispuesto en el código civil y de forma análoga el RLCE.")
    parrafo_agregado = buscar_add_parrafo3(document, parrafo_agregado, "")
    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"La penalidad por mora se calcula de acuerdo a la siguiente formula:")
    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"")

    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"PENALIDAD DIARIA = (0.10*M)/(F*P)")
    parrafo_agregado = buscar_add_parrafo3(document, parrafo_agregado, "")
    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"M: Monto Vigente")

    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"F: 0.40 para plazos menores o iguales a 60 días F: 0.25 para plazos mayores a 60 días.")

    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"P: Plazo vigente en días.")
    parrafo_agregado = buscar_add_parrafo3(document, parrafo_agregado, "")
    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"La acumulación del monto máximo de la penalidad y/o de otras penalidades en la ejecución del servicio dará lugar a la resolución de la orden de servicio.")
    parrafo_agregado = buscar_add_parrafo3(document, parrafo_agregado, "")
    parrafo_agregado=buscar_add_parrafo3(document,parrafo_agregado,"Si el contratista incurre en retraso injustificado a la presentación de cada entregable, se le aplicará una penalidad por cada día de atraso, en concordancia con los artículos 161, 162 y 163 del Reglamento de la Ley de Contrataciones del Estado.")

    # 5. Titulo OTRAS CONDICIONES
    replace_parrafo(document, 'OBSERVANCIA DEL CÓDIGO DE ÉTICA', "OTRAS CONDICIONES")
    buscar_add_parrafo2(document, 'RESPONSABILIDAD DE VICIOS OCULTOS',
                        "De no cumplir con lo estipulado en los párrafos precedentes se resolverá la Orden de Servicio de forma unilateral.")

    # Guardando el Documento
    document.save(rutaTdrscambiados + x[len(x) - 1]);
