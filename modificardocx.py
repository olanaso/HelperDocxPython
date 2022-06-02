#agregando al pie de pagina
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
import shutil
files_source="C:\\Users\\ASUS\\Downloads\\tdrcuscoapu\\APURIMAC\\"
files_destino="C:\\Users\\ASUS\\Downloads\\tdrcuscoapu\\APURIMAC\\procesado\\"



files = [os.path.join(r,file) for r,d,f in os.walk(files_source) for file in f]
arr_txt = [x for x in files if x.endswith(".docx")]

for file in arr_txt:
    x = file.split("\\") #archivos separados
    document = Document(file)
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(7)
    for section in document.sections:
        footer = section.footer
        print(footer.paragraphs)
        paragraph=footer.add_paragraph("\t\t"+x[len(x)-1])
        paragraph.style = document.styles['Normal']
        paragraph.alignment = 2
    print('---CORRECTO---'+file)
    document.save(files_destino+x[len(x)-1])
