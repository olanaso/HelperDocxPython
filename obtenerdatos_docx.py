# coding=utf-8
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
import os
import pandas as pd
import mpu.io
import csv
import re

def buscar_posicion_add_parrafo(doc, textSearchParaffo):
    indice=0
    for i,p in enumerate(doc.paragraphs):
        #print(i, p.text)
        if textSearchParaffo.decode('utf-8') in p.text:
            paragraph=p
            indice=i
    return indice


rutaTdrs="C:\\Users\\ASUS\\Downloads\\TDRS\\"
rutaTdrscambiados="D:\\GTS\TDRS\\abril\\test\\pro\\"

#obteniendo los archivos word de cada uno
files = [os.path.join(r,file) for r,d,f in os.walk(rutaTdrs) for file in f]
arr_txt = [x for x in files if x.endswith(".docx")]
print(len(arr_txt))


data = []

for i in range(0, len(arr_txt)):
    print('Item-------------------------------------:'+str(i))
    print(arr_txt[i])
    document = Document(arr_txt[i])

    posicion1=buscar_posicion_add_parrafo(document, "ALCANCES Y DESCRIPCIÓN DEL SERVICIO")
    #print (posicion1)
    posicion2 = buscar_posicion_add_parrafo(document, "DURACION DEL SERVICIO") or buscar_posicion_add_parrafo(document, "DURACIÓN DEL SERVICIO")

    posicion3 = buscar_posicion_add_parrafo(document, "OBJETO DE LA CONTRATACIÓN")


    #print (posicion2)
    arrayactividades=[]
    tablaarrayactividades = ["N°\tDescripción de Servicio (Actividades que se han cumplido)\tComentario"]
    contadortabla=1
    for j in range(posicion1+1,posicion2):
        #print(document.paragraphs[i].text)

        if (len(document.paragraphs[j].text.strip()) == 0):
            print('sin act')
        else:
            arrayactividades.append(document.paragraphs[j].text)
            tablaarrayactividades.append((str(contadortabla)+'\t'+document.paragraphs[j].text+'\tSi cumple').encode('utf-8'))
            contadortabla=contadortabla+1

    #print '--------------'
    text_actividades='\n'.join(arrayactividades)
    text_tablaactividades='\n'.join(tablaarrayactividades)
    #print text_actividades
    # print '--------------'

    print arr_txt[i]
    arraauxcod=arr_txt[i].split('\\')
    codigo=arraauxcod[len(arraauxcod)-1]
    objetoTDR=''
    print len(document.paragraphs[posicion3+1].text.strip())

    if(len(document.paragraphs[posicion3+1].text.strip())==0):

        objetoTDR=document.paragraphs[posicion3 + 2].text
    else:
        objetoTDR = document.paragraphs[posicion3 + 1].text

    data.append({'codigo':codigo.split('.')[-2], 'actividades':text_actividades,'tabla_actividades':text_tablaactividades.decode('utf-8'), 'objeto':objetoTDR})



df = pd.DataFrame.from_dict(data)

df.to_excel('act2.xlsx')
